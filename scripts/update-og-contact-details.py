from pathlib import Path

from docx import Document


ROOT = Path(r"C:\OGBuild\OG Products\OG Web.site")
DOCS = [
    ROOT / "output" / "doc" / "og-letterhead-template.docx",
    ROOT / "output" / "doc" / "og-quote-template.docx",
    ROOT / "output" / "doc" / "og-invoice-template.docx",
]

REPLACEMENTS = {
    "hello@og-web.site  |  www.og-web.site": "og@og-web.site  |  www.og-web.site  |  0786615047",
    "hello@og-web.site": "og@og-web.site",
    "[Phone / direct contact]": "0786615047",
    "[Phone]": "0786615047",
}


def replace_in_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        for old, new in REPLACEMENTS.items():
            if old in run.text:
                run.text = run.text.replace(old, new)


def replace_in_table(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph)
            for nested in cell.tables:
                replace_in_table(nested)


def update_document(path: Path) -> None:
    doc = Document(path)
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    for table in doc.tables:
        replace_in_table(table)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                replace_in_paragraph(paragraph)
            for table in part.tables:
                replace_in_table(table)
    doc.save(path)


for doc_path in DOCS:
    update_document(doc_path)
    print(doc_path)
