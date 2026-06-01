from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\OGBuild\OG Products\OG Web.site")
OUTPUT_DIR = ROOT / "output" / "doc"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

LOGO = ROOT / "public" / "assets" / "branding" / "Untitled-3.png"


BLACK = "0D0F0F"
CARBON = "171A1B"
GREEN = "8CFF41"
GREEN_DARK = "74D630"
GREEN_PALE = "EAF8DE"
GREEN_MIST = "F4FBEF"
WHITE = "FFFFFF"
SILVER = "CDD3D6"
SLATE = "5E676D"
ASH = "EEF1F3"
MID = "D6DDE1"
TEXT = "1B1F23"


def set_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.36)
    section.left_margin = Inches(0.52)
    section.right_margin = Inches(0.52)
    section.header_distance = Inches(0.14)
    section.footer_distance = Inches(0.14)
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(10)
    doc.core_properties.author = "OpenAI Codex"
    doc.core_properties.company = "OG Web.site"
    doc.core_properties.comments = "Branded business document template."
    background = OxmlElement("w:background")
    background.set(qn("w:color"), WHITE)
    doc._element.insert(0, background)


def usable_width(doc: Document) -> Inches:
    section = doc.sections[0]
    return section.page_width - section.left_margin - section.right_margin


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge_data in edges.items():
        edge = tc_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_borders.append(edge)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                edge.set(qn(f"w:{key}"), str(edge_data[key]))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def remove_cell_borders(cell) -> None:
    set_cell_border(
        cell,
        left={"val": "nil"},
        top={"val": "nil"},
        right={"val": "nil"},
        bottom={"val": "nil"},
    )


def set_table_widths(table, widths) -> None:
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = width


def clear_paragraph_spacing(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_together = True


def add_text(
    paragraph,
    text: str,
    *,
    font="Segoe UI",
    size=10.5,
    bold=False,
    italic=False,
    color=TEXT,
    caps=False,
) -> None:
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.all_caps = caps


def add_band(doc: Document, fill: str, label: str, *, label_color=GREEN, after=2) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = usable_width(doc)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    remove_cell_borders(cell)
    set_cell_margins(cell, top=22, start=100, bottom=22, end=100)
    p = cell.paragraphs[0]
    clear_paragraph_spacing(p)
    add_text(p, label, font="Bahnschrift", size=8, bold=True, color=label_color, caps=True)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph().paragraph_format.space_after = Pt(after)


def add_logo(paragraph, width: float) -> None:
    run = paragraph.add_run()
    run.add_picture(str(LOGO), width=Inches(width))


def add_meta_card(cell, label: str, value: str) -> None:
    shade_cell(cell, WHITE)
    set_cell_margins(cell, top=60, start=90, bottom=60, end=90)
    set_cell_border(
        cell,
        top={"val": "single", "sz": 18, "color": GREEN},
        left={"val": "single", "sz": 8, "color": MID},
        right={"val": "single", "sz": 8, "color": MID},
        bottom={"val": "single", "sz": 8, "color": MID},
    )
    label_p = cell.paragraphs[0]
    clear_paragraph_spacing(label_p)
    add_text(label_p, label, font="Bahnschrift", size=7.2, bold=True, color=SLATE, caps=True)
    value_p = cell.add_paragraph()
    clear_paragraph_spacing(value_p)
    add_text(value_p, value, font="Bahnschrift", size=10.5, bold=True, color=BLACK)


def add_party_block(cell, label: str, lines, *, fill=WHITE) -> None:
    shade_cell(cell, fill)
    set_cell_margins(cell, top=65, start=105, bottom=70, end=105)
    set_cell_border(
        cell,
        top={"val": "single", "sz": 14, "color": GREEN},
        left={"val": "single", "sz": 8, "color": MID},
        right={"val": "single", "sz": 8, "color": MID},
        bottom={"val": "single", "sz": 8, "color": MID},
    )
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        clear_paragraph_spacing(p)
        if idx == 0:
            add_text(p, label, font="Bahnschrift", size=7.2, bold=True, color=SLATE, caps=True)
        elif idx == 1:
            add_text(p, line, font="Bahnschrift", size=11.8, bold=True, color=BLACK)
        else:
            add_text(p, line, size=9.3, color=TEXT)


def add_footer(section, left_text: str, center_text: str, right_text: str) -> None:
    footer = section.footer
    table = footer.add_table(rows=1, cols=3, width=section.page_width - section.left_margin - section.right_margin)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [Inches(2.4), Inches(2.1), Inches(2.57)])
    for cell in table.rows[0].cells:
        shade_cell(cell, BLACK)
        remove_cell_borders(cell)
        set_cell_margins(cell, top=38, start=90, bottom=38, end=90)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    values = [left_text, center_text, right_text]
    aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT]
    for cell, value, align in zip(table.rows[0].cells, values, aligns):
        p = cell.paragraphs[0]
        clear_paragraph_spacing(p)
        p.alignment = align
        add_text(p, value, size=7.4, color=WHITE)


def add_table_header(row, labels, *, fill=BLACK, color=GREEN) -> None:
    for cell, label in zip(row.cells, labels):
        shade_cell(cell, fill)
        set_cell_margins(cell, top=45, start=70, bottom=45, end=70)
        set_cell_border(
            cell,
            left={"val": "single", "sz": 8, "color": BLACK},
            right={"val": "single", "sz": 8, "color": BLACK},
            top={"val": "single", "sz": 8, "color": BLACK},
            bottom={"val": "single", "sz": 8, "color": BLACK},
        )
        p = cell.paragraphs[0]
        clear_paragraph_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, label, font="Bahnschrift", size=7.4, bold=True, color=color, caps=True)


def fill_body_cell(cell, text: str, *, align=WD_ALIGN_PARAGRAPH.LEFT, fill=WHITE, bold=False, color=TEXT) -> None:
    shade_cell(cell, fill)
    set_cell_margins(cell, top=38, start=65, bottom=38, end=65)
    set_cell_border(
        cell,
        left={"val": "single", "sz": 8, "color": MID},
        right={"val": "single", "sz": 8, "color": MID},
        top={"val": "single", "sz": 8, "color": MID},
        bottom={"val": "single", "sz": 8, "color": MID},
    )
    p = cell.paragraphs[0]
    clear_paragraph_spacing(p)
    p.alignment = align
    add_text(p, text, size=9.2, bold=bold, color=color)


def add_invoice() -> Path:
    doc = Document()
    set_page(doc)
    doc.core_properties.title = "OG Invoice Template"

    add_band(doc, BLACK, "Design. Build. Hosting.")

    hero = doc.add_table(rows=1, cols=2)
    hero.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(hero, [Inches(5.02), Inches(1.9)])

    left = hero.cell(0, 0)
    right = hero.cell(0, 1)

    shade_cell(left, WHITE)
    set_cell_margins(left, top=36, start=48, bottom=32, end=48)
    set_cell_border(
        left,
        top={"val": "single", "sz": 10, "color": MID},
        left={"val": "single", "sz": 10, "color": MID},
        bottom={"val": "single", "sz": 10, "color": MID},
        right={"val": "single", "sz": 10, "color": GREEN},
    )
    logo_p = left.paragraphs[0]
    clear_paragraph_spacing(logo_p)
    logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_logo(logo_p, 2.32)
    tag = left.add_paragraph()
    clear_paragraph_spacing(tag)
    tag.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(tag, "CLIENT INVOICE TEMPLATE", font="Bahnschrift", size=7.2, bold=True, color=GREEN, caps=True)

    shade_cell(right, BLACK)
    set_cell_margins(right, top=44, start=84, bottom=44, end=84)
    remove_cell_borders(right)
    title = right.paragraphs[0]
    clear_paragraph_spacing(title)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(title, "INVOICE", font="Bahnschrift", size=16.8, bold=True, color=WHITE, caps=True)
    sub = right.add_paragraph()
    clear_paragraph_spacing(sub)
    add_text(sub, "Sharper presentation.", size=8.1, color=SILVER)
    sub = right.add_paragraph()
    clear_paragraph_spacing(sub)
    add_text(sub, "Cleaner payment flow.", size=8.1, color=SILVER)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)

    meta = doc.add_table(rows=1, cols=4)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(meta, [Inches(1.76), Inches(1.76), Inches(1.76), Inches(1.79)])
    add_meta_card(meta.cell(0, 0), "Invoice no.", "[INV-2026-001]")
    add_meta_card(meta.cell(0, 1), "Issue date", "[28 Apr 2026]")
    add_meta_card(meta.cell(0, 2), "Due date", "[05 May 2026]")
    add_meta_card(meta.cell(0, 3), "Terms", "[7 days]")

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    parties = doc.add_table(rows=1, cols=2)
    parties.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(parties, [Inches(3.54), Inches(3.53)])
    add_party_block(
        parties.cell(0, 0),
        "Issued by",
        [
            "label",
            "OG Web.site",
            "og@og-web.site",
            "www.og-web.site",
            "0786615047",
            "[Business address]",
        ],
    )
    add_party_block(
        parties.cell(0, 1),
        "Bill to",
        [
            "label",
            "[Client / business name]",
            "[Contact person]",
            "[Client email]",
            "[Client phone]",
            "[Billing address]",
        ],
        fill=GREEN_MIST,
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    item_band = doc.add_paragraph()
    clear_paragraph_spacing(item_band)
    add_text(item_band, "Invoice breakdown", font="Bahnschrift", size=9.2, bold=True, color=BLACK, caps=True)

    items = doc.add_table(rows=3, cols=5)
    items.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(items, [Inches(3.3), Inches(0.75), Inches(1.0), Inches(1.0), Inches(1.0)])
    add_table_header(items.rows[0], ["Description", "Qty", "Rate", "VAT", "Total"])
    rows = [
        ("[Line item description]", "[1]", "[500.00]", "[100.00]", "[600.00]"),
        ("[Second line item]", "[1]", "[350.00]", "[70.00]", "[420.00]"),
    ]
    fills = [WHITE, GREEN_MIST]
    for row_index, values in enumerate(rows, start=1):
        for cell, value in zip(items.rows[row_index].cells, values):
            align = WD_ALIGN_PARAGRAPH.LEFT if cell == items.rows[row_index].cells[0] else WD_ALIGN_PARAGRAPH.CENTER
            fill_body_cell(cell, value, align=align, fill=fills[row_index - 1])

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    summary_wrap = doc.add_table(rows=1, cols=2)
    summary_wrap.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(summary_wrap, [Inches(4.62), Inches(2.44)])

    left_note = summary_wrap.cell(0, 0)
    shade_cell(left_note, GREEN_PALE)
    set_cell_margins(left_note, top=95, start=120, bottom=95, end=120)
    set_cell_border(
        left_note,
        top={"val": "single", "sz": 8, "color": MID},
        left={"val": "single", "sz": 8, "color": MID},
        right={"val": "single", "sz": 8, "color": MID},
        bottom={"val": "single", "sz": 8, "color": MID},
    )
    p = left_note.paragraphs[0]
    clear_paragraph_spacing(p)
    add_text(p, "Project note", font="Bahnschrift", size=8.2, bold=True, color=SLATE, caps=True)
    p = left_note.add_paragraph()
    clear_paragraph_spacing(p)
    add_text(
        p,
        "Milestones, ownership notes, or a short summary of what the client is paying for.",
        size=9.2,
        color=TEXT,
    )

    totals = summary_wrap.cell(0, 1)
    shade_cell(totals, WHITE)
    set_cell_border(
        totals,
        top={"val": "single", "sz": 8, "color": MID},
        left={"val": "single", "sz": 8, "color": MID},
        right={"val": "single", "sz": 8, "color": MID},
        bottom={"val": "single", "sz": 8, "color": MID},
    )
    set_cell_margins(totals, top=55, start=0, bottom=55, end=0)
    totals_table = totals.add_table(rows=4, cols=2)
    totals_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(totals_table, [Inches(1.15), Inches(1.13)])
    total_rows = [
        ("Subtotal", "[850.00]", WHITE, TEXT),
        ("VAT", "[170.00]", GREEN_MIST, TEXT),
        ("Deposit paid", "[-0.00]", WHITE, TEXT),
        ("Total due", "[1,020.00]", GREEN, BLACK),
    ]
    for idx, (label, value, fill, text_color) in enumerate(total_rows):
        fill_body_cell(totals_table.cell(idx, 0), label, fill=fill, bold=idx == 3, color=text_color)
        fill_body_cell(
            totals_table.cell(idx, 1),
            value,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
            fill=fill,
            bold=True,
            color=text_color,
        )

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    payment = doc.add_table(rows=1, cols=2)
    payment.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(payment, [Inches(2.2), Inches(4.87)])

    pay_left = payment.cell(0, 0)
    shade_cell(pay_left, BLACK)
    remove_cell_borders(pay_left)
    set_cell_margins(pay_left, top=62, start=96, bottom=62, end=96)
    p = pay_left.paragraphs[0]
    clear_paragraph_spacing(p)
    add_text(p, "Payment details", font="Bahnschrift", size=8.6, bold=True, color=GREEN, caps=True)
    for line in (
        "[Bank name]",
        "[Account name]",
        "[Sort code]",
        "[Account number]",
        "Ref: [Invoice no.]",
    ):
        p = pay_left.add_paragraph()
        clear_paragraph_spacing(p)
        add_text(p, line, size=8.8, color=WHITE)

    pay_right = payment.cell(0, 1)
    shade_cell(pay_right, WHITE)
    set_cell_margins(pay_right, top=68, start=110, bottom=68, end=110)
    set_cell_border(
        pay_right,
        top={"val": "single", "sz": 8, "color": MID},
        left={"val": "single", "sz": 8, "color": MID},
        right={"val": "single", "sz": 8, "color": MID},
        bottom={"val": "single", "sz": 8, "color": MID},
    )
    p = pay_right.paragraphs[0]
    clear_paragraph_spacing(p)
    add_text(p, "Terms", font="Bahnschrift", size=7.2, bold=True, color=SLATE, caps=True)
    for text in (
        "Payment due within [7] days unless otherwise agreed in writing.",
        "Late fees or milestone rules can be added here.",
        "Use the invoice number as the transfer reference.",
    ):
        p = pay_right.add_paragraph()
        clear_paragraph_spacing(p)
        add_text(p, text, size=8.9, color=TEXT)

    add_footer(
        doc.sections[0],
        "OG Web.site",
        "og@og-web.site  |  www.og-web.site  |  0786615047",
        "Replace placeholders before sending",
    )

    path = OUTPUT_DIR / "og-invoice-template.docx"
    doc.save(path)
    return path


def add_quote() -> Path:
    doc = Document()
    set_page(doc)
    doc.core_properties.title = "OG Quote Template"

    add_band(doc, GREEN, "Quote / proposal", label_color=BLACK)

    hero = doc.add_table(rows=1, cols=2)
    hero.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(hero, [Inches(2.15), Inches(4.77)])

    left = hero.cell(0, 0)
    right = hero.cell(0, 1)

    shade_cell(left, BLACK)
    remove_cell_borders(left)
    set_cell_margins(left, top=48, start=90, bottom=48, end=90)
    p = left.paragraphs[0]
    clear_paragraph_spacing(p)
    add_text(p, "QUOTE", font="Bahnschrift", size=16.2, bold=True, color=WHITE, caps=True)
    p = left.add_paragraph()
    clear_paragraph_spacing(p)
    add_text(p, "Website design / build / hosting", size=8.0, color=SILVER)

    shade_cell(right, WHITE)
    set_cell_border(
        right,
        top={"val": "single", "sz": 10, "color": MID},
        left={"val": "single", "sz": 10, "color": GREEN},
        bottom={"val": "single", "sz": 10, "color": MID},
        right={"val": "single", "sz": 10, "color": MID},
    )
    set_cell_margins(right, top=44, start=56, bottom=44, end=56)
    p = right.paragraphs[0]
    clear_paragraph_spacing(p)
    add_logo(p, 2.34)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    meta = doc.add_table(rows=1, cols=4)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(meta, [Inches(2.0), Inches(1.55), Inches(1.5), Inches(1.87)])
    add_meta_card(meta.cell(0, 0), "Prepared for", "[Client / brand]")
    add_meta_card(meta.cell(0, 1), "Quote no.", "[Q-2026-001]")
    add_meta_card(meta.cell(0, 2), "Date", "[28 Apr 2026]")
    add_meta_card(meta.cell(0, 3), "Valid until", "[12 May 2026]")

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    overview = doc.add_table(rows=1, cols=2)
    overview.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(overview, [Inches(5.0), Inches(2.07)])
    left = overview.cell(0, 0)
    right = overview.cell(0, 1)
    add_party_block(
        left,
        "Project overview",
        [
            "label",
            "[Project title / workstream]",
            "Explain what is being built, what problem it solves, and what the client is approving.",
            "Keep it direct and tight.",
        ],
    )
    shade_cell(right, GREEN_MIST)
    set_cell_margins(right, top=95, start=130, bottom=95, end=130)
    set_cell_border(
        right,
        top={"val": "single", "sz": 14, "color": GREEN},
        left={"val": "single", "sz": 8, "color": MID},
        right={"val": "single", "sz": 8, "color": MID},
        bottom={"val": "single", "sz": 8, "color": MID},
    )
    p = right.paragraphs[0]
    clear_paragraph_spacing(p)
    add_text(p, "Prepared by", font="Bahnschrift", size=8, bold=True, color=SLATE, caps=True)
    for text in (
        "OG Web.site",
        "og@og-web.site",
        "www.og-web.site",
        "0786615047",
    ):
        p = right.add_paragraph()
        clear_paragraph_spacing(p)
        add_text(
            p,
            text,
            font="Bahnschrift" if text == "OG Web.site" else "Segoe UI",
            size=11.2 if text == "OG Web.site" else 8.9,
            bold=text == "OG Web.site",
            color=BLACK if text == "OG Web.site" else TEXT,
        )

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    clear_paragraph_spacing(p)
    add_text(p, "Deliverables and investment", font="Bahnschrift", size=9.2, bold=True, color=BLACK, caps=True)

    items = doc.add_table(rows=5, cols=4)
    items.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(items, [Inches(1.35), Inches(3.47), Inches(1.15), Inches(1.1)])
    add_table_header(items.rows[0], ["Stage", "Deliverable", "Timeline", "Fee"])
    rows = [
        ("01", "Discovery and first-screen direction", "[2 days]", "[350.00]"),
        ("02", "Multi-page design and build", "[5 days]", "[1,250.00]"),
        ("03", "Contact form, domain, email and launch setup", "[2 days]", "[400.00]"),
        ("04", "Post-launch edits / aftercare block", "[30 days]", "[250.00]"),
    ]
    fills = [WHITE, GREEN_MIST, WHITE, GREEN_MIST]
    for row_index, values in enumerate(rows, start=1):
        for cell, value in zip(items.rows[row_index].cells, values):
            align = WD_ALIGN_PARAGRAPH.CENTER if cell != items.rows[row_index].cells[1] else WD_ALIGN_PARAGRAPH.LEFT
            fill_body_cell(cell, value, align=align, fill=fills[row_index - 1], bold=cell == items.rows[row_index].cells[0])

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    timeline = doc.add_table(rows=1, cols=3)
    timeline.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(timeline, [Inches(2.32), Inches(2.32), Inches(2.11)])
    stages = [
        ("01", "Deposit / approval", "Quote signed and deposit received before work starts."),
        ("02", "Build window", "Design, revisions and launch prep handled in one block."),
        ("03", "Go live", "Final balance due on launch or on the agreed milestone date."),
    ]
    for cell, (number, heading, text) in zip(timeline.rows[0].cells, stages):
        shade_cell(cell, WHITE)
        set_cell_margins(cell, top=85, start=110, bottom=90, end=110)
        set_cell_border(
            cell,
            top={"val": "single", "sz": 14, "color": GREEN},
            left={"val": "single", "sz": 8, "color": MID},
            right={"val": "single", "sz": 8, "color": MID},
            bottom={"val": "single", "sz": 8, "color": MID},
        )
        p = cell.paragraphs[0]
        clear_paragraph_spacing(p)
        add_text(p, number, font="Bahnschrift", size=9.4, bold=True, color=GREEN_DARK, caps=True)
        p = cell.add_paragraph()
        clear_paragraph_spacing(p)
        add_text(p, heading, font="Bahnschrift", size=10.2, bold=True, color=BLACK)
        p = cell.add_paragraph()
        clear_paragraph_spacing(p)
        add_text(p, text, size=8.7, color=TEXT)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    wrap = doc.add_table(rows=1, cols=2)
    wrap.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(wrap, [Inches(4.25), Inches(2.5)])
    left = wrap.cell(0, 0)
    right = wrap.cell(0, 1)

    shade_cell(left, BLACK)
    remove_cell_borders(left)
    set_cell_margins(left, top=90, start=120, bottom=90, end=120)
    p = left.paragraphs[0]
    clear_paragraph_spacing(p)
    add_text(p, "Commercial notes", font="Bahnschrift", size=8.8, bold=True, color=GREEN, caps=True)
    for text in (
        "Deposit required: [50%]",
        "Quote valid until: [Date]",
        "Revisions included: [2 rounds]",
        "Out-of-scope billed separately at [rate]",
    ):
        p = left.add_paragraph()
        clear_paragraph_spacing(p)
        add_text(p, text, size=8.9, color=WHITE)

    shade_cell(right, GREEN_PALE)
    set_cell_margins(right, top=90, start=120, bottom=90, end=120)
    set_cell_border(
        right,
        top={"val": "single", "sz": 14, "color": GREEN},
        left={"val": "single", "sz": 8, "color": MID},
        right={"val": "single", "sz": 8, "color": MID},
        bottom={"val": "single", "sz": 8, "color": MID},
    )
    p = right.paragraphs[0]
    clear_paragraph_spacing(p)
    add_text(p, "Total investment", font="Bahnschrift", size=7.2, bold=True, color=SLATE, caps=True)
    p = right.add_paragraph()
    clear_paragraph_spacing(p)
    add_text(p, "[2,250.00]", font="Bahnschrift", size=17.6, bold=True, color=BLACK)
    p = right.add_paragraph()
    clear_paragraph_spacing(p)
    add_text(p, "Acceptance", font="Bahnschrift", size=7.2, bold=True, color=SLATE, caps=True)
    for text in ("Name: ____________________", "Signature: ____________________", "Date: ____________________"):
        p = right.add_paragraph()
        clear_paragraph_spacing(p)
        add_text(p, text, size=9, color=TEXT)

    add_footer(
        doc.sections[0],
        "OG Web.site",
        "Quote template for branded web projects",
        "Replace pricing, dates and scope before sending",
    )

    path = OUTPUT_DIR / "og-quote-template.docx"
    doc.save(path)
    return path


def add_letterhead() -> Path:
    doc = Document()
    set_page(doc)
    doc.core_properties.title = "OG Letterhead Template"

    section = doc.sections[0]
    header = section.header
    head_top = header.add_table(rows=2, cols=2, width=usable_width(doc))
    head_top.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(head_top, [Inches(4.55), Inches(2.52)])

    for cell in head_top.rows[0].cells:
        shade_cell(cell, BLACK)
        remove_cell_borders(cell)
        set_cell_margins(cell, top=30, start=80, bottom=30, end=80)
    head_top.rows[0].cells[0].merge(head_top.rows[0].cells[1])
    p = head_top.rows[0].cells[0].paragraphs[0]
    clear_paragraph_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(p, "Design. Build. Hosting.", font="Bahnschrift", size=8.2, bold=True, color=GREEN, caps=True)

    logo_cell = head_top.rows[1].cells[0]
    info_cell = head_top.rows[1].cells[1]
    for cell in (logo_cell, info_cell):
        shade_cell(cell, WHITE)
        set_cell_border(
            cell,
            bottom={"val": "single", "sz": 16, "color": GREEN},
            left={"val": "nil"},
            right={"val": "nil"},
            top={"val": "nil"},
        )
        set_cell_margins(cell, top=80, start=80, bottom=110, end=80)
    p = logo_cell.paragraphs[0]
    clear_paragraph_spacing(p)
    add_logo(p, 3.1)
    p = info_cell.paragraphs[0]
    clear_paragraph_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(p, "OG Web.site", font="Bahnschrift", size=13, bold=True, color=BLACK)
    for text in (
            "og@og-web.site",
        "www.og-web.site",
            "0786615047",
        "[Address / company details]",
    ):
        p = info_cell.add_paragraph()
        clear_paragraph_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_text(p, text, size=9.6, color=TEXT)

    add_footer(
        section,
        "OG Web.site",
        "og@og-web.site  |  www.og-web.site  |  0786615047",
        "[Address]  |  [Company number / VAT number]",
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph()
    clear_paragraph_spacing(p)
    add_text(p, "[Date]", font="Bahnschrift", size=10.5, bold=True, color=BLACK)

    p = doc.add_paragraph()
    clear_paragraph_spacing(p)
    add_text(p, "[Recipient name]", font="Bahnschrift", size=12, bold=True, color=BLACK)
    for line in ("[Company / organisation]", "[Address line 1]", "[Address line 2]", "[Postcode]"):
        p = doc.add_paragraph()
        clear_paragraph_spacing(p)
        add_text(p, line, size=10.4, color=TEXT)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    subject = doc.add_table(rows=1, cols=1)
    subject.alignment = WD_TABLE_ALIGNMENT.CENTER
    subject.autofit = False
    subject.columns[0].width = usable_width(doc)
    cell = subject.cell(0, 0)
    shade_cell(cell, BLACK)
    remove_cell_borders(cell)
    set_cell_margins(cell, top=90, start=140, bottom=90, end=140)
    p = cell.paragraphs[0]
    clear_paragraph_spacing(p)
    add_text(p, "Subject: [Letter subject / project topic]", font="Bahnschrift", size=11, bold=True, color=GREEN)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    paragraphs = [
        "Dear [Recipient name],",
        "This file is the branded OG letterhead template. Replace the placeholders, keep the subject line tight, and type directly over the sample copy with the actual letter content.",
        "Use the body for proposals, client updates, approvals, project handover notes, or any formal communication that needs to look cleaner than a plain Word page.",
        "If you need a second page, keep writing - the header and footer stay in place so the branding carries through automatically.",
        "Kind regards,",
        "[Your name]",
        "OG Web.site",
    ]
    for text in paragraphs:
        p = doc.add_paragraph()
        clear_paragraph_spacing(p)
        p.paragraph_format.space_after = Pt(8)
        add_text(
            p,
            text,
            font="Bahnschrift" if text in {"Dear [Recipient name],", "Kind regards,", "[Your name]", "OG Web.site"} else "Segoe UI",
            size=10.8 if text != "OG Web.site" else 11.2,
            bold=text in {"Dear [Recipient name],", "[Your name]", "OG Web.site"},
            color=BLACK if text in {"Dear [Recipient name],", "[Your name]", "OG Web.site"} else TEXT,
        )

    path = OUTPUT_DIR / "og-letterhead-template.docx"
    doc.save(path)
    return path


def main() -> None:
    add_invoice()
    add_quote()
    add_letterhead()


if __name__ == "__main__":
    main()
