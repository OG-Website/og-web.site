from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\OGBuild\OG Products\OG Web.site")
OUTPUT_DIR = ROOT / "output" / "doc"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
LOGO = ROOT / "public" / "assets" / "branding" / "Untitled-3.png"
SUPPLIED_CROPPED_LOGO = ROOT / "public" / "assets" / "branding" / "logo cropped.png"
OPAQUE_LOGO = ROOT / "tmp" / "docs" / "og-logo-white-header.png"
CROPPED_LOGO = ROOT / "tmp" / "docs" / "og-logo-white-header-cropped.png"

GREEN = "8CFF41"
GREEN_DEEP = "3F7F22"
GREEN_PALE = "F4FBEF"
GREEN_LINE = "BDEFA2"
TEXT = "24302A"
MUTED = "66736B"
LINE = "D8E2DA"
SOFT = "F7FAF7"
WHITE = "FFFFFF"


def set_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.12)
    section.footer_distance = Inches(0.16)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    doc.core_properties.author = "OG Web.site"
    doc.core_properties.company = "OG Web.site"
    doc.core_properties.title = "Website Review and Content Audit Blank Template"


def prepare_opaque_logo() -> None:
    if SUPPLIED_CROPPED_LOGO.exists():
        CROPPED_LOGO.parent.mkdir(parents=True, exist_ok=True)
        CROPPED_LOGO.write_bytes(SUPPLIED_CROPPED_LOGO.read_bytes())
        return
    if not LOGO.exists():
        return
    OPAQUE_LOGO.parent.mkdir(parents=True, exist_ok=True)
    try:
        from PIL import Image

        image = Image.open(LOGO).convert("RGBA")
        white = Image.new("RGBA", image.size, (255, 255, 255, 255))
        white.alpha_composite(image)
        rgb = white.convert("RGB")
        rgb.save(OPAQUE_LOGO)

        pixels = rgb.load()
        width, height = rgb.size
        threshold = 246
        min_x, min_y = width, height
        max_x, max_y = 0, 0
        for y in range(height):
            for x in range(width):
                r, g, b = pixels[x, y]
                if not (r >= threshold and g >= threshold and b >= threshold):
                    min_x = min(min_x, x)
                    min_y = min(min_y, y)
                    max_x = max(max_x, x)
                    max_y = max(max_y, y)
        if min_x < max_x and min_y < max_y:
            pad = 8
            crop_box = (
                max(min_x - pad, 0),
                max(min_y - pad, 0),
                min(max_x + pad, width),
                min(max_y + pad, height),
            )
            rgb.crop(crop_box).save(CROPPED_LOGO)
        else:
            rgb.save(CROPPED_LOGO)
    except Exception:
        OPAQUE_LOGO.write_bytes(LOGO.read_bytes())
        CROPPED_LOGO.write_bytes(LOGO.read_bytes())
    if not CROPPED_LOGO.exists() and OPAQUE_LOGO.exists():
        CROPPED_LOGO.write_bytes(OPAQUE_LOGO.read_bytes())


def usable_width(doc: Document):
    section = doc.sections[0]
    return section.page_width - section.left_margin - section.right_margin


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LINE, size=8, top=True, right=True, bottom=True, left=True) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, enabled in (("top", top), ("right", right), ("bottom", bottom), ("left", left)):
        edge = tc_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_borders.append(edge)
        edge.set(qn("w:val"), "single" if enabled else "nil")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths) -> None:
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = width


def clear_spacing(paragraph, after=0) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.0


def add_text(paragraph, text, *, size=9.5, bold=False, color=TEXT, font="Aptos", caps=False) -> None:
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.all_caps = caps


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    f = footer.add_table(rows=1, cols=3, width=usable_width(doc))
    f.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(f, [Inches(2.2), Inches(2.7), Inches(2.2)])
    for cell in f.rows[0].cells:
        shade_cell(cell, WHITE)
        set_cell_border(cell, color=GREEN_LINE, size=8, top=True, right=False, left=False, bottom=False)
        set_cell_margins(cell, top=35, start=20, bottom=20, end=20)
    footer_values = ["OG Web.site", "Website Review & Content Audit", "Blank proofing template"]
    footer_aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT]
    for cell, value, align in zip(f.rows[0].cells, footer_values, footer_aligns):
        p = cell.paragraphs[0]
        clear_spacing(p)
        p.alignment = align
        add_text(p, value, size=7.5, color=MUTED)


def add_body_header(doc: Document) -> None:
    h = doc.add_table(rows=1, cols=2)
    h.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(h, [Inches(4.55), Inches(2.55)])
    for cell in h.rows[0].cells:
        shade_cell(cell, WHITE)
        set_cell_border(cell, color=GREEN, size=14, top=False, right=False, left=False, bottom=True)
        set_cell_margins(cell, top=16, start=25, bottom=24, end=25)
    p = h.cell(0, 0).paragraphs[0]
    clear_spacing(p)
    if CROPPED_LOGO.exists():
        p.add_run().add_picture(str(CROPPED_LOGO), width=Inches(1.85))
    elif OPAQUE_LOGO.exists():
        p.add_run().add_picture(str(OPAQUE_LOGO), width=Inches(1.85))
    elif LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(1.85))
    else:
        add_text(p, "OG Web.site", size=18, bold=True, color=GREEN_DEEP, font="Bahnschrift")
    p = h.cell(0, 1).paragraphs[0]
    clear_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(p, "OG Web.site", size=11, bold=True, color=TEXT, font="Bahnschrift")
    for line in ("og@og-web.site", "www.og-web.site", "0786615047"):
        p = h.cell(0, 1).add_paragraph()
        clear_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_text(p, line, size=7.8, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(5)


def add_title(doc: Document, title: str, subtitle: str = "") -> None:
    p = doc.add_paragraph()
    clear_spacing(p, after=2)
    add_text(p, title, size=18, bold=True, color=TEXT, font="Aptos Display", caps=True)
    if subtitle:
        p = doc.add_paragraph()
        clear_spacing(p, after=10)
        add_text(p, subtitle, size=9.5, color=MUTED)


def add_cover(doc: Document) -> None:
    add_body_header(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    add_title(
        doc,
        "Website Review & Content Audit Workbook",
        "Blank branded template for reviewing website pages, content, images, mobile/desktop experience and final actions.",
    )

    details = doc.add_table(rows=7, cols=2)
    details.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(details, [Inches(2.0), Inches(5.1)])
    rows = [
        ("Website / Project", ""),
        ("Client / Prepared For", ""),
        ("Review Date", ""),
        ("Version", ""),
        ("Prepared By", "OG Web.site"),
        ("Email", "og@og-web.site"),
        ("Phone", "0786615047"),
    ]
    for index, (label, value) in enumerate(rows):
        left, right = details.rows[index].cells
        shade_cell(left, GREEN_PALE)
        shade_cell(right, WHITE)
        for cell in (left, right):
            set_cell_border(cell)
            set_cell_margins(cell, top=85, start=90, bottom=85, end=90)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = left.paragraphs[0]
        clear_spacing(p)
        add_text(p, label, size=8.2, bold=True, color=GREEN_DEEP, caps=True)
        p = right.paragraphs[0]
        clear_spacing(p)
        add_text(p, value, size=9.5, color=TEXT)

    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    note = doc.add_table(rows=1, cols=1)
    note.alignment = WD_TABLE_ALIGNMENT.CENTER
    note.columns[0].width = usable_width(doc)
    cell = note.cell(0, 0)
    shade_cell(cell, SOFT)
    set_cell_border(cell, color=GREEN_LINE, size=10)
    set_cell_margins(cell, top=120, start=120, bottom=120, end=120)
    p = cell.paragraphs[0]
    clear_spacing(p)
    add_text(p, "Purpose", size=8.2, bold=True, color=GREEN_DEEP, caps=True)
    p = cell.add_paragraph()
    clear_spacing(p)
    add_text(
        p,
        "Use this document to collect clear page-by-page feedback before website changes are agreed and completed.",
        size=10,
        color=TEXT,
    )
    doc.add_page_break()


def add_instructions(doc: Document) -> None:
    add_body_header(doc)
    add_title(doc, "Instructions", "Keep the workbook blank until the actual website pages and screenshots are added.")
    items = [
        "Complete one review sheet per website page.",
        "Paste screenshots into the screenshot box or attach them separately.",
        "Use the text change box only for wording that needs replacing.",
        "Mark each page as High, Medium or Low priority.",
        "Use the final action plan to turn feedback into agreed tasks.",
    ]
    table = doc.add_table(rows=len(items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [Inches(0.5), Inches(6.6)])
    for index, item in enumerate(items, start=1):
        num, text = table.rows[index - 1].cells
        shade_cell(num, GREEN_PALE)
        shade_cell(text, WHITE if index % 2 else SOFT)
        for cell in (num, text):
            set_cell_border(cell)
            set_cell_margins(cell, top=70, start=70, bottom=70, end=70)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = num.paragraphs[0]
        clear_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, str(index), size=9, bold=True, color=GREEN_DEEP)
        p = text.paragraphs[0]
        clear_spacing(p)
        add_text(p, item, size=9.5)
    doc.add_page_break()


def add_line_box(cell, title: str, lines=3, fill=WHITE) -> None:
    shade_cell(cell, fill)
    set_cell_border(cell)
    set_cell_margins(cell, top=55, start=70, bottom=55, end=70)
    p = cell.paragraphs[0]
    clear_spacing(p, after=3)
    add_text(p, title, size=7.7, bold=True, color=GREEN_DEEP, caps=True)
    for _ in range(lines):
        p = cell.add_paragraph()
        clear_spacing(p, after=5)
        add_text(p, "_" * 48, size=8, color=LINE)


def add_page_sheet(doc: Document, number: int) -> None:
    add_body_header(doc)
    add_title(doc, "Page Review Sheet", f"Blank page review sheet {number}.")

    top = doc.add_table(rows=2, cols=4)
    top.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(top, [Inches(1.25), Inches(2.35), Inches(1.25), Inches(2.25)])
    labels = ["Page Name", "", "Page URL", "", "Priority", "[ ] High   [ ] Medium   [ ] Low", "Screenshot", "[ ] Attached / pasted below"]
    for cell, value in zip([c for r in top.rows for c in r.cells], labels):
        shade_cell(cell, GREEN_PALE if value in {"Page Name", "Page URL", "Priority", "Screenshot"} else WHITE)
        set_cell_border(cell)
        set_cell_margins(cell, top=55, start=70, bottom=55, end=70)
        p = cell.paragraphs[0]
        clear_spacing(p)
        add_text(
            p,
            value,
            size=7.8 if value in {"Page Name", "Page URL", "Priority", "Screenshot"} else 9,
            bold=value in {"Page Name", "Page URL", "Priority", "Screenshot"},
            color=GREEN_DEEP if value in {"Page Name", "Page URL", "Priority", "Screenshot"} else TEXT,
            caps=value in {"Page Name", "Page URL", "Priority", "Screenshot"},
        )

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    shot = doc.add_table(rows=1, cols=1)
    shot.alignment = WD_TABLE_ALIGNMENT.CENTER
    shot.columns[0].width = usable_width(doc)
    cell = shot.cell(0, 0)
    shade_cell(cell, SOFT)
    set_cell_border(cell, color=GREEN_LINE, size=10)
    set_cell_margins(cell, top=170, start=100, bottom=170, end=100)
    p = cell.paragraphs[0]
    clear_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "Paste screenshot here", size=10, color=MUTED)

    fields = [
        "First Impression",
        "What I Like",
        "What I Don't Like",
        "Text Changes Required",
        "Font & Typography Review",
        "Colours & Branding Review",
        "Images Review",
        "Backgrounds Review",
        "Mobile Phone Review",
        "Desktop Review",
        "SEO Improvements",
        "Missing Content",
        "New Images Required",
        "Notes",
    ]
    grid = doc.add_table(rows=7, cols=2)
    grid.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(grid, [Inches(3.55), Inches(3.55)])
    for index, label in enumerate(fields):
        add_line_box(grid.rows[index // 2].cells[index % 2], label, lines=2, fill=WHITE if index % 4 in (0, 1) else SOFT)
    doc.add_page_break()


def add_website_wide(doc: Document) -> None:
    add_body_header(doc)
    add_title(doc, "Website-Wide Branding Review", "Use this section for comments that apply across more than one page.")
    fields = ["Logo Use", "Brand Consistency", "Tone Of Voice", "Typography", "Colours", "Navigation", "Calls To Action", "Overall Notes"]
    grid = doc.add_table(rows=4, cols=2)
    grid.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(grid, [Inches(3.55), Inches(3.55)])
    for index, label in enumerate(fields):
        add_line_box(grid.rows[index // 2].cells[index % 2], label, lines=4, fill=WHITE if index % 4 in (0, 1) else SOFT)
    doc.add_page_break()


def add_table_page(doc: Document, title: str, subtitle: str, headers: list[str], rows: int) -> None:
    add_body_header(doc)
    add_title(doc, title, subtitle)
    table = doc.add_table(rows=rows + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    width = usable_width(doc) / len(headers)
    set_table_widths(table, [width for _ in headers])
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, GREEN_PALE)
        set_cell_border(cell, color=GREEN_LINE, size=10)
        set_cell_margins(cell, top=55, start=50, bottom=55, end=50)
        p = cell.paragraphs[0]
        clear_spacing(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, header, size=7.6, bold=True, color=GREEN_DEEP, caps=True)
    for row_index in range(1, rows + 1):
        for cell in table.rows[row_index].cells:
            shade_cell(cell, WHITE if row_index % 2 else SOFT)
            set_cell_border(cell)
            set_cell_margins(cell, top=80, start=50, bottom=80, end=50)
            p = cell.paragraphs[0]
            clear_spacing(p)
            add_text(p, " ", size=9)
    doc.add_page_break()


def build() -> Path:
    prepare_opaque_logo()
    doc = Document()
    set_page(doc)
    add_footer(doc)
    add_cover(doc)
    add_instructions(doc)
    for number in range(1, 9):
        add_page_sheet(doc, number)
    add_website_wide(doc)
    add_table_page(
        doc,
        "Image Requirements",
        "List any images that need to be sourced, created, replaced, cropped or resized.",
        ["Page", "Image Needed", "Purpose", "Priority", "Notes"],
        10,
    )
    add_table_page(
        doc,
        "Final Action Plan",
        "Turn the review into clear tasks before sign-off.",
        ["Action Required", "Page", "Owner", "Priority", "Status"],
        12,
    )
    add_table_page(
        doc,
        "Priority Matrix",
        "Define what High, Medium and Low mean for this project.",
        ["Priority", "Meaning", "Examples"],
        5,
    )
    path = OUTPUT_DIR / "og-website-proofing-template.docx"
    try:
        doc.save(path)
    except PermissionError:
        path = OUTPUT_DIR / "og-website-proofing-template-corrected.docx"
        doc.save(path)
    return path


if __name__ == "__main__":
    print(build())
